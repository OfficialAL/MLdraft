from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

output = 'Final Assignment Report\\MLOPS_Final_Report_Draft.docx'

authors = ['Alex Rush','Juzer Pakawala','Michael Wagg','Rishabh Panikar']

doc = Document()

# Title page
p = doc.add_paragraph()
run = p.add_run('MLOPS Final Assignment - Winrate Prediction')
run.bold = True
run.font.size = Pt(20)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()
sub = doc.add_paragraph()
sub_run = sub.add_run('Data Management and Machine Learning Operations')
sub_run.italic = True
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()
authors_p = doc.add_paragraph('Authors: ' + ', '.join(authors))
authors_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_page_break()

# Simple Table of Contents (manual)
doc.add_heading('Table of Contents', level=1)
sections = ['Abstract','Introduction','Data Sources','Data Processing','Feature Engineering','Modeling','Evaluation and Results','Deployment and MLOps','Discussion','Conclusion','Contributions','References']
for i, s in enumerate(sections,1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_page_break()

# Sections with placeholder / repo-derived text
doc.add_heading('Abstract', level=1)
doc.add_paragraph('This report summarises the project to collect, process and model League of Legends match data to predict team winrate. The work covers data ingestion, feature engineering, model training, evaluation, and considerations for deploying the model within an MLOps workflow.')

doc.add_heading('Introduction', level=1)
doc.add_paragraph('The objective of this project is to explore the use of game match data to predict win probabilities and to demonstrate reproducible MLOps practices including data versioning, preprocessing pipelines, model training and basic deployment planning.')

doc.add_heading('Data Sources', level=1)
doc.add_paragraph('Primary dataset: League of Legends match data collected via the Riot API. Additional candidate datasets considered included F1 telemetry data. The chosen dataset provides player roles, champion picks, match outcomes and timestamps.')

doc.add_heading('Data Processing', level=1)
doc.add_paragraph('Raw match JSON is parsed and processed into feature matrices. Scripts are in `Winrate_Prediction/src` and `scripts` (e.g., `fetch_data.py`, `prepare_features.py`). Missing values, role assignments and champion mappings are handled in preprocessing steps.')

doc.add_heading('Feature Engineering', level=1)
doc.add_paragraph('Features include role-normalised champion statistics, aggregated team-level features, and meta-features capturing recent performance. Feature caching and matrix building scripts are provided in `scripts/build_and_cache_matrices.py` and `Winrate_Prediction/src/feature_engineering.py`.')

doc.add_heading('Modeling', level=1)
doc.add_paragraph('Models trained include gradient boosting and logistic-type classifiers. Training orchestration and experiment scripts are under `Winrate_Prediction/src/train_model.py`. Metrics tracked: accuracy, ROC-AUC and per-role performance.')

doc.add_heading('Evaluation and Results', level=1)
doc.add_paragraph('Key results are retained in `models/metrics.json` and `Winrate_Prediction/analysis_outputs`. The report will include tables and figures exported from notebooks and analysis outputs. (Placeholders included in this draft.)')

doc.add_heading('Deployment and MLOps', level=1)
doc.add_paragraph('Deployment considerations: model packaging, versioning, CI for training pipelines, reproducible environments via `requirements.txt` and virtualenv, and logging/monitoring. Recommendation: wrap feature pipeline and model in a container and serve via a lightweight API with automated retraining triggers.')

doc.add_heading('Discussion', level=1)
doc.add_paragraph('Discuss challenges such as data drift, role inference mismatches, and balancing evaluation across rare champions. Outline next steps for improving dataset coverage and automated evaluation.')

doc.add_heading('Conclusion', level=1)
doc.add_paragraph('This project demonstrates a working pipeline from data collection to model evaluation and an MLOps-aware plan for deployment. Future work includes production-grade monitoring and scheduled retraining.')

doc.add_heading('Contributions', level=1)
for a in authors:
    doc.add_paragraph(f'{a}: [list contributions here]')

doc.add_heading('References', level=1)
doc.add_paragraph('Relevant notebooks and scripts are in the repository. Cite APIs and libraries used (Riot API, scikit-learn, XGBoost, python-docx).')

# Save

doc.save(output)
print('Wrote', output)
