"""Convert a specific Markdown file (v2) to DOCX using python-docx.

Run from repository root with the project's venv Python.
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path

md_path = Path('..') / 'Final Assignment Report' / 'MLOPS_Project_Condensed_v2.md'
docx_path = Path('..') / 'Final Assignment Report' / 'MLOPS_Project_Condensed_v2.docx'

doc = Document()

in_code = False
code_lines = []

with md_path.open(encoding='utf-8') as f:
    for raw in f:
        line = raw.rstrip('\n')
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                in_code = False
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        stripped = line.lstrip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
            continue

        if line.strip().startswith('---'):
            doc.add_page_break()
            continue

        if line.strip() == '':
            doc.add_paragraph('')
            continue

        doc.add_paragraph(line)

docx_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(docx_path)
print(f'Wrote {docx_path.resolve()}')
