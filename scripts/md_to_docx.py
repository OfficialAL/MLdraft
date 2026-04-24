"""Convert a Markdown file to a simple DOCX preserving headings, lists, and code blocks.

Usage: run with the project's virtualenv Python from repository root.
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path

md_path = Path('..') / 'Final Assignment Report' / 'MLOPS_Project_Condensed.md'
docx_path = Path('..') / 'Final Assignment Report' / 'MLOPS_Project_Condensed.docx'

doc = Document()

in_code = False
code_lines = []

with md_path.open(encoding='utf-8') as f:
    for raw in f:
        line = raw.rstrip('\n')
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # flush code block
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                in_code = False
            continue
        if in_code:
            code_lines.append(line)
            continue

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        # Lists (simple bullets)
        stripped = line.lstrip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
            continue

        # Horizontal rule
        if line.strip().startswith('---'):
            doc.add_page_break()
            continue

        # Normal paragraph (skip if empty)
        if line.strip() == '':
            # add blank paragraph for spacing
            doc.add_paragraph('')
            continue

        # Inline backticks -> leave as-is
        doc.add_paragraph(line)

# Save
docx_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(docx_path)
print(f'Wrote {docx_path.resolve()}')
